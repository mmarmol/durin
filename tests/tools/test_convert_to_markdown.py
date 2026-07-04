"""Tests for the convert_to_markdown tool.

Fixtures are real documents generated on the fly (python-docx for .docx,
plain text for .html/.csv) so the tests exercise the actual markitdown
conversion path. PDF is intentionally not covered here — a fabricated
PDF byte stream would test garbage, not the converter; PDF is covered by
live verification.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from durin.agent.tools.convert_to_markdown import ConvertToMarkdownTool


def _make_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading("Alpha Report", level=1)
    document.add_paragraph("Opening paragraph of the report.")
    document.add_heading("Findings", level=2)
    document.add_paragraph("Detail line about the findings.")
    document.save(str(path))


@pytest.mark.asyncio
async def test_docx_conversion_preserves_heading_hierarchy(tmp_path: Path) -> None:
    source = tmp_path / "report.docx"
    _make_docx(source)
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path=str(source))

    assert "error" not in result
    assert result["format"] == ".docx"
    assert "# Alpha Report" in result["markdown"]
    assert "## Findings" in result["markdown"]
    assert "Opening paragraph of the report." in result["markdown"]
    assert result["outline"] == ["# Alpha Report", "## Findings"]
    assert result["size_chars"] == len(result["markdown"])


@pytest.mark.asyncio
async def test_html_conversion_preserves_heading_hierarchy(tmp_path: Path) -> None:
    source = tmp_path / "page.html"
    source.write_text(
        "<html><body><h1>Main Title</h1><p>Intro text.</p>"
        "<h2>Subsection</h2><p>Body text.</p></body></html>",
        encoding="utf-8",
    )
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path=str(source))

    assert "error" not in result
    assert result["format"] == ".html"
    assert "# Main Title" in result["markdown"]
    assert "## Subsection" in result["markdown"]
    assert result["outline"] == ["# Main Title", "## Subsection"]


@pytest.mark.asyncio
async def test_csv_converts_to_table(tmp_path: Path) -> None:
    source = tmp_path / "data.csv"
    source.write_text("name,age\nalice,30\nbob,25\n", encoding="utf-8")
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path=str(source))

    assert "error" not in result
    assert result["format"] == ".csv"
    assert "|" in result["markdown"]
    assert "alice" in result["markdown"]
    assert "bob" in result["markdown"]


@pytest.mark.asyncio
async def test_missing_file_returns_error(tmp_path: Path) -> None:
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path=str(tmp_path / "nope.docx"))

    assert "error" in result
    assert "not found" in result["error"]


@pytest.mark.asyncio
async def test_unsupported_format_names_it_and_lists_supported(
    tmp_path: Path,
) -> None:
    source = tmp_path / "doc.odt"
    source.write_bytes(b"not a real odt")
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path=str(source))

    assert "error" in result
    assert ".odt" in result["error"]
    assert ".docx" in result["error"]
    assert ".pdf" in result["error"]


@pytest.mark.asyncio
async def test_empty_path_returns_error(tmp_path: Path) -> None:
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path="   ")

    assert result == {"error": "path is required"}


@pytest.mark.asyncio
async def test_relative_path_resolves_against_workspace(tmp_path: Path) -> None:
    source = tmp_path / "page.html"
    source.write_text("<h1>Rel</h1>", encoding="utf-8")
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path="page.html")

    assert "error" not in result
    assert result["path"] == str(source)
    assert "# Rel" in result["markdown"]


@pytest.mark.asyncio
async def test_result_key_order_puts_markdown_last(tmp_path: Path) -> None:
    source = tmp_path / "page.html"
    source.write_text("<h1>Order</h1><p>body</p>", encoding="utf-8")
    tool = ConvertToMarkdownTool(workspace=tmp_path)

    result = await tool.execute(path=str(source))

    assert list(result.keys()) == [
        "path",
        "format",
        "size_chars",
        "outline",
        "markdown",
    ]


def test_tool_is_read_only(tmp_path: Path) -> None:
    tool = ConvertToMarkdownTool(workspace=tmp_path)
    assert tool.read_only is True
    assert tool.name == "convert_to_markdown"
