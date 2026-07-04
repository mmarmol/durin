"""Phase 2a: document conversion at ingest + structural chunking.

The additive half of Phase 2 — supported documents are converted to markdown
at ingest, kept verbatim alongside a ``source.md`` sidecar, and chunked
structurally with heading breadcrumbs. The Library recall-scope flip is a
separate change; these tests do not assert recall isolation.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from durin.memory.doc_convert import (
    DocConvertError,
    convert_file_to_markdown,
    is_convertible,
)
from durin.memory.ingestion import ingest_artifact
from durin.memory.reference import (
    chunk_structured,
    ingest_reference,
    reference_chunks,
)
from durin.utils.helpers import estimate_text_tokens


def _make_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading("Alpha Report", level=1)
    document.add_paragraph("Opening paragraph of the report.")
    document.add_heading("Findings", level=2)
    document.add_paragraph("Detail line about the findings.")
    document.save(str(path))


# --- shared converter helper -------------------------------------------------


def test_convert_docx_to_markdown(tmp_path: Path) -> None:
    src = tmp_path / "report.docx"
    _make_docx(src)
    out = convert_file_to_markdown(src)
    assert out.suffix == ".docx"
    assert "# Alpha Report" in out.markdown
    assert "## Findings" in out.markdown


def test_convert_unsupported_format_raises(tmp_path: Path) -> None:
    src = tmp_path / "doc.odt"
    src.write_bytes(b"not a real odt")
    with pytest.raises(DocConvertError) as exc:
        convert_file_to_markdown(src)
    assert ".odt" in str(exc.value)


def test_is_convertible() -> None:
    assert is_convertible(".PDF")
    assert is_convertible(".docx")
    assert is_convertible(".epub")
    assert not is_convertible(".md")
    assert not is_convertible(".txt")


# --- structural chunker ------------------------------------------------------


def test_chunk_structured_breadcrumb_from_heading_stack() -> None:
    md = (
        "# Book\n\nintro paragraph.\n\n"
        "## Chapter One\n\nbody of chapter one.\n\n"
        "### Details\n\nnested detail body.\n"
    )
    recs = chunk_structured(md)
    crumbs = {r["breadcrumb"] for r in recs}
    assert "Book" in crumbs
    assert "Book › Chapter One" in crumbs
    assert "Book › Chapter One › Details" in crumbs
    # Heading lines are captured in the breadcrumb, never duplicated into text.
    assert all("#" not in r["text"] for r in recs)


def test_chunk_structured_no_headings_empty_breadcrumb() -> None:
    recs = chunk_structured("just a plain paragraph, no headings at all.")
    assert len(recs) == 1
    assert recs[0]["breadcrumb"] == ""


def test_chunk_structured_size_capped() -> None:
    big = "## Section\n\n" + ("This is a sentence. " * 500)
    recs = chunk_structured(big, max_tokens=384)
    assert len(recs) > 1
    assert all(estimate_text_tokens(r["text"]) <= 384 for r in recs)
    assert all(r["breadcrumb"] == "Section" for r in recs)


# --- ingest converts internally ----------------------------------------------


def test_ingest_docx_converts_keeps_original_and_sidecar(tmp_path: Path) -> None:
    ws = tmp_path / "ws"
    ws.mkdir()
    src = tmp_path / "report.docx"
    _make_docx(src)

    result = ingest_artifact(ws, src)
    # The returned content is the CONVERTED markdown, not raw docx bytes.
    assert "# Alpha Report" in result["content"]
    assert "## Findings" in result["content"]

    entry_dir = Path(result["source"]).parent
    assert (entry_dir / "source.docx").exists()  # verbatim original kept
    assert (entry_dir / "source.md").exists()  # markdown rendering sidecar


def test_ingest_docx_idempotent_over_original_bytes(tmp_path: Path) -> None:
    ws = tmp_path / "ws"
    ws.mkdir()
    src = tmp_path / "report.docx"
    _make_docx(src)

    first = ingest_artifact(ws, src)
    second = ingest_artifact(ws, src)
    assert second["id"] == first["id"]


def test_ingest_unsupported_binary_rejected(tmp_path: Path) -> None:
    from durin.memory.ingestion import IngestError

    ws = tmp_path / "ws"
    ws.mkdir()
    src = tmp_path / "image.png"
    src.write_bytes(b"\x89PNG\r\n\x1a\n not really a png")
    with pytest.raises(IngestError):
        ingest_artifact(ws, src)


def test_reference_chunks_carry_breadcrumb_and_parent(tmp_path: Path) -> None:
    ws = tmp_path / "ws"
    ws.mkdir()
    md = "# Title\n\nintro.\n\n## Part A\n\nbody of part a.\n"
    res = ingest_reference(ws, "mydoc", md)

    recs = reference_chunks(ws, res.ref)
    assert recs
    for r in recs:
        assert set(("idx", "parent", "tokens", "text", "breadcrumb")) <= set(r)
        assert r["parent"] == res.ref
        assert r["tokens"] <= 384
    assert any(r["breadcrumb"] == "Title › Part A" for r in recs)
